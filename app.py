import os
from pathlib import Path
from flask import Flask, render_template, request, send_file, flash, redirect
import pdfplumber
import docx
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from werkzeug.utils import secure_filename
import google.generativeai as genai
from fpdf import FPDF
from concurrent.futures import ThreadPoolExecutor
from docx import Document
from datetime import datetime
import logging
import pytesseract
from PIL import Image

# ============================
# 🔑 API Key Configuration
# ============================
os.environ["GOOGLE_API_KEY"] = "AIzaSyCGFQvsNk7LpRQklcp5k4481wdc-eZ_qGw"
genai.configure(api_key=os.environ["GOOGLE_API_KEY"])
model = genai.GenerativeModel("models/gemini-1.5-pro")

# ============================
# ⚙ Flask App Configuration
# ============================
app = Flask(__name__)
app.secret_key = "supersecretkey"  # Required for flash messages
UPLOAD_FOLDER, RESULTS_FOLDER = Path("uploads"), Path("results")
ALLOWED_EXTENSIONS = {"pdf", "txt", "docx"}

app.config.update(UPLOAD_FOLDER=UPLOAD_FOLDER, RESULTS_FOLDER=RESULTS_FOLDER)
executor = ThreadPoolExecutor(max_workers=4)

UPLOAD_FOLDER.mkdir(exist_ok=True)
RESULTS_FOLDER.mkdir(exist_ok=True)


# ============================
# 📂 Helper Functions
# ============================
def allowed_file(filename):
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXTENSIONS


def extract_text_from_file(file_path):
    ext = file_path.suffix.lower()
    try:
        if ext == ".pdf":
            with pdfplumber.open(file_path) as pdf:
                return "\n".join(page.extract_text() or "" for page in pdf.pages)
        elif ext == ".docx":
            return "\n".join(para.text for para in docx.Document(file_path).paragraphs)
        elif ext == ".txt":
            return file_path.read_text(encoding="utf-8")
    except Exception as e:
        logging.error(f"Error extracting text: {str(e)}")
        return None
    return ""


def analyze_image(image_path):
    """Extract text from an image using Tesseract OCR."""
    try:
        # Set the path to the Tesseract executable (if not in PATH)
        pytesseract.pytesseract.tesseract_cmd = (
            r"C:\Program Files\Tesseract-OCR\tesseract.exe"  # Windows
        )
        # pytesseract.pytesseract.tesseract_cmd = '/usr/local/bin/tesseract'  # macOS/Linux

        # Open the image using PIL
        img = Image.open(image_path)
        # Use pytesseract to extract text
        text = pytesseract.image_to_string(img)
        return text.strip() if text else ""
    except Exception as e:
        logging.error(f"Error analyzing image: {str(e)}")
        return ""


def extract_text_and_images_from_pdf(file_path):
    """Extract text and images from a PDF file."""
    text = ""
    images = []
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            text += page.extract_text() or ""
            for image in page.images:
                # Save the image temporarily
                image_path = RESULTS_FOLDER / f"image_{len(images)}.png"
                with open(image_path, "wb") as img_file:
                    img_file.write(image["stream"].get_data())
                images.append(image_path)
    return text, images


# Rest of the code remains the same...
# ============================
# 📂 Helper Functions
# ============================


def generate_mcqs_from_text_and_images(text, images, num_questions, difficulty):
    """Generate MCQs from text and images using Gemini API."""
    # Analyze images and extract text
    image_texts = []
    for image in images:
        image_text = analyze_image(image)
        image_texts.append(image_text)

    # Combine text and image content
    full_content = text + "\n".join(image_texts)

    # Define difficulty-specific instructions
    difficulty_instructions = {
        "easy": """
        - Generate simple and straightforward questions suitable for beginners.
        - Focus on basic concepts and facts.
        - Avoid complex or ambiguous wording.
        """,
        "medium": """
        - Generate moderately challenging questions suitable for intermediate learners.
        - Include questions that require understanding of relationships between concepts.
        - Use some application-based or scenario-based questions.
        """,
        "hard": """
        - Generate complex and challenging questions suitable for advanced learners.
        - Include questions that require critical thinking, analysis, or synthesis.
        - Use application-based, scenario-based, or problem-solving questions.
        """,
    }

    prompt = f"""
    You are an expert MCQ generator. Your task is to generate {num_questions} multiple-choice questions (MCQs) based on the following text and any diagrams or visual content:

    ### Instructions:
    - Difficulty Level: {difficulty.capitalize()}
    {difficulty_instructions[difficulty]}
    - Ensure questions are **clear, relevant, and unambiguous**.
    - Provide **4 distinct answer choices** for each question.
    - Indicate the **correct answer** at the end of each question.
    - Carefully analyze both the text and any diagrams or visual content to ensure the questions and distractors are accurate and relevant.
    -And don't generate question number 

    ### Text Content:
    '{full_content}'

    ### Output Format:
    For each MCQ, use the following format:
    ## MCQ
    Question: [Your question here]
    A) [Option A]
    B) [Option B]
    C) [Option C]
    D) [Option D]
    Correct Answer: [Correct option letter]

    

    Now, generate {num_questions} MCQs based on the provided text and difficulty level.
    """
    try:
        response = model.generate_content(prompt)
        return response.text.strip() if response else "No MCQs generated."
    except Exception as e:
        logging.error(f"Error generating MCQs: {str(e)}")
        return None


def save_text_file(content, filename):
    """Save the generated MCQs to a text file."""
    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    file_path = RESULTS_FOLDER / filename
    content_with_timestamp = f"Generated on: {timestamp}\n\n{content}"
    try:
        file_path.write_text(content_with_timestamp, encoding="utf-8")
        return file_path
    except Exception as e:
        logging.error(f"Error saving text file: {str(e)}")
        return None


from fpdf import FPDF
from datetime import datetime
import os
import logging


def generate_pdf(mcqs, filename):
    """Generate a PDF file from the MCQs."""
    try:
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=12)  # Use default font (Arial or Helvetica)
        pdf.cell(0, 10, "Generated MCQs", ln=True, align="C")
        pdf.set_font("Arial", size=12)

        # Remove any carriage return characters
        mcqs = mcqs.replace("\r", "").strip()

        for i, mcq in enumerate(
            mcqs.split("## MCQ"), start=0
        ):  # Add question numbering
            if mcq.strip():
                lines = mcq.strip().split("\n")
                correct_answer = None

                # Extract the correct answer
                for line in lines:
                    if "Correct Answer:" in line:
                        correct_answer = line.split(":")[-1].strip()
                        break

                # Extract question text safely
                question_text = (
                    lines[0].split("Question:", 1)[-1].strip()
                    if "Question:" in lines[0]
                    else lines[0].strip()
                )

                # Add question number and question text (bolded)
                pdf.set_font("Arial", "B", 12)  # Bold for question
                pdf.multi_cell(
                    0,
                    10,
                    f"Q{i}: {question_text}".encode("latin-1", "replace").decode(
                        "latin-1"
                    ),
                )
                pdf.ln(2)  # Small gap after question

                # Add options
                pdf.set_font("Arial", size=12)  # Normal font for options
                for line in lines[1:]:
                    if "Correct Answer:" in line:
                        continue  # Skip the correct answer line

                    if line.strip():  # Ensure the line is not empty
                        # Remove numbers like "1)", "2)", "3)", "4)" from the options
                        option_text = line.strip()
                        if option_text[:2].isdigit() and option_text[2] == ")":
                            option_text = option_text[
                                3:
                            ].strip()  # Remove the number and parenthesis

                        # Highlight correct answer in green
                        if correct_answer and line.strip().startswith(
                            correct_answer + ")"
                        ):
                            pdf.set_text_color(0, 128, 0)  # Green for correct answer
                            pdf.set_font("Arial", "B", 12)  # Bold for correct answer
                        else:
                            pdf.set_text_color(0, 0, 0)  # Black for incorrect options
                            pdf.set_font(
                                "Arial", size=12
                            )  # Normal font for incorrect options

                        # Wrap long text into multiple lines
                        pdf.multi_cell(
                            0,
                            10,
                            option_text.encode("latin-1", "replace").decode("latin-1"),
                        )
                        pdf.ln(2)  # Small gap between options

                # Add correct answer at the end
                pdf.set_text_color(0, 128, 0)  # Green for correct answer
                pdf.set_font("Arial", "B", 12)  # Bold for correct answer
                pdf.multi_cell(
                    0,
                    10,
                    f"Correct Answer: {correct_answer}".encode(
                        "latin-1", "replace"
                    ).decode("latin-1"),
                )
                pdf.ln(10)  # Spacing between MCQs

                # Reset text color
                pdf.set_text_color(0, 0, 0)

        # Add timestamp
        pdf.set_font("Arial", size=10)
        pdf.cell(
            0,
            10,
            f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
            ln=True,
        )

        # Save the PDF file
        pdf_path = os.path.join(app.config["RESULTS_FOLDER"], filename)
        pdf.output(pdf_path)
        return pdf_path

    except Exception as e:
        logging.error(f"Error generating PDF: {str(e)}")
        return None


def generate_docx(mcqs, filename):
    """Generate a DOCX file from the MCQs."""
    try:
        doc = Document()

        # Add a title and timestamp
        title = doc.add_paragraph("Generated MCQs")
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        title.runs[0].bold = True
        title.runs[0].font.size = Pt(14)

        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        doc.add_paragraph(f"Generated on: {timestamp}\n").alignment = (
            WD_PARAGRAPH_ALIGNMENT.CENTER
        )

        for i, mcq in enumerate(
            mcqs.split("## MCQ"), start=0
        ):  # Add question numbering
            if mcq.strip():
                lines = mcq.strip().split("\n")
                correct_answer = None

                # Extract the correct answer
                for line in lines:
                    if "Correct Answer:" in line:
                        correct_answer = line.split(":")[-1].strip()
                        break

                # Extract question text safely
                question_text = (
                    lines[0].split("Question:", 1)[-1].strip()
                    if "Question:" in lines[0]
                    else lines[0].strip()
                )

                # Add question number and question text (bolded)
                question_para = doc.add_paragraph()
                question_para.add_run(f"Q{i}: {question_text}").bold = True

                # Add options
                for line in lines[1:]:
                    if "Correct Answer:" in line:
                        continue  # Skip the correct answer line

                    if line.strip():  # Ensure the line is not empty
                        # Remove numbers like "1)", "2)", "3)", "4)" from the options
                        option_text = line.strip()
                        if option_text[:2].isdigit() and option_text[2] == ")":
                            option_text = option_text[
                                3:
                            ].strip()  # Remove the number and parenthesis

                        option_para = doc.add_paragraph(option_text)
                        # Highlight the correct answer in green
                        if correct_answer and line.strip().startswith(
                            correct_answer + ")"
                        ):
                            for run in option_para.runs:
                                run.font.color.rgb = RGBColor(
                                    0, 128, 0
                                )  # Green for correct answer
                                run.bold = True  # Bold for correct answer

                # Add correct answer at the end
                correct_para = doc.add_paragraph(f"Correct Answer: {correct_answer}")
                for run in correct_para.runs:
                    run.font.color.rgb = RGBColor(0, 128, 0)  # Green for correct answer
                    run.bold = True  # Bold for correct answer

                doc.add_paragraph("\n")  # Add a line break between questions

        # Add timestamp
        timestamp_para = doc.add_paragraph(
            f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
        )
        timestamp_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        docx_path = os.path.join(app.config["RESULTS_FOLDER"], filename)
        doc.save(docx_path)
        return docx_path
    except Exception as e:
        logging.error(f"Error generating DOCX: {str(e)}")
        return None


# ============================
# 📌 Flask Routes
# ============================
@app.route("/")
def index():
    return render_template("index.html")


@app.route("/about")
def about():
    return render_template("about.html")


@app.route("/contact")
def contact():
    return render_template("contact.html")


@app.route("/generate", methods=["POST"])
def generate_mcqs():
    if "file" not in request.files:
        flash("No file uploaded", "error")
        return redirect(request.url)

    file = request.files["file"]
    if file.filename == "":
        flash("No file selected", "error")
        return redirect(request.url)

    if not allowed_file(file.filename):
        flash("Invalid file format. Allowed formats: PDF, TXT, DOCX", "error")
        return redirect(request.url)

    filename = secure_filename(file.filename)
    file_path = UPLOAD_FOLDER / filename
    file.save(file_path)

    if file_path.suffix.lower() == ".pdf":
        text, images = extract_text_and_images_from_pdf(file_path)
    else:
        text = extract_text_from_file(file_path)
        images = []

    if not text:
        flash("Could not extract text from file", "error")
        return redirect(request.url)

    try:
        num_questions = int(request.form["num_questions"])
        if num_questions <= 0:
            flash("Number of questions must be greater than zero", "error")
            return redirect(request.url)
    except ValueError:
        flash("Invalid number of questions", "error")
        return redirect(request.url)

    # Get the selected difficulty level
    difficulty = request.form["difficulty"]

    # Generate MCQs based on the difficulty level
    mcqs = generate_mcqs_from_text_and_images(text, images, num_questions, difficulty)
    if not mcqs:
        flash("Failed to generate MCQs", "error")
        return redirect(request.url)

    base_filename = filename.rsplit(".", 1)[0]
    txt_filename = f"generated_mcqs_{base_filename}.txt"
    pdf_filename = f"generated_mcqs_{base_filename}.pdf"
    docx_filename = f"generated_mcqs_{base_filename}.docx"

    if not save_text_file(mcqs, txt_filename):
        flash("Failed to save text file", "error")
        return redirect(request.url)

    if not generate_pdf(mcqs, pdf_filename):
        flash("Failed to generate PDF", "error")
        return redirect(request.url)

    if not generate_docx(mcqs, docx_filename):
        flash("Failed to generate DOCX", "error")
        return redirect(request.url)

    return render_template(
        "results.html",
        mcqs=mcqs,
        txt_filename=txt_filename,
        pdf_filename=pdf_filename,
        docx_filename=docx_filename,
    )


@app.route("/download/<filename>")
def download_file(filename):
    file_path = RESULTS_FOLDER / filename
    if not file_path.exists():
        flash("File not found", "error")
        return redirect("/")
    return send_file(file_path, as_attachment=True)


# ============================
# 🚀 Run the Flask App
# ============================
if __name__ == "__main__":
    logging.basicConfig(level=logging.INFO)
    app.run(host="0.0.0.0", port=5000, debug=True, threaded=True)
