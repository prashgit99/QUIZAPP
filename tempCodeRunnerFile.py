import os
from pathlib import Path
from flask import Flask, render_template, request, send_file, flash, redirect
import pdfplumber
import docx
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from werkzeug.utils import secure_filename
import google.generativeai as genai
from fpdf import FPDF
from concurrent.futures import ThreadPoolExecutor
from docx import Document
from datetime import datetime
import logging

# ============================
# 🔑 API Key Configuration
# ============================
os.environ["GOOGLE_API_KEY"] = "AIzaSyAs0ngUkfdmnIjRowyfsJ26-aGqx8c5Xl4"
genai.configure(api_key=os.environ["GOOGLE_API_KEY"])
model = genai.GenerativeModel("models/gemini-1.5-pro")

# ============================
# ⚙ Flask App Configuration
# ============================
app = Flask(__name__)
app.secret_key = "supersecretkey"  # Required for flash messages
UPLOAD_FOLDER, RESULTS_FOLDER = Path("uploads"), Path("results")
ALLOWED_EXTENSIONS = {'pdf', 'txt', 'docx'}

app.config.update(UPLOAD_FOLDER=UPLOAD_FOLDER, RESULTS_FOLDER=RESULTS_FOLDER)
executor = ThreadPoolExecutor(max_workers=4)

UPLOAD_FOLDER.mkdir(exist_ok=True)
RESULTS_FOLDER.mkdir(exist_ok=True)

# ============================
# 📂 Helper Functions
# ============================
def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def extract_text_from_file(file_path):
    ext = file_path.suffix.lower()
    try:
        if ext == '.pdf':
            with pdfplumber.open(file_path) as pdf:
                return "\n".join(page.extract_text() or '' for page in pdf.pages)
        elif ext == '.docx':
            return "\n".join(para.text for para in docx.Document(file_path).paragraphs)
        elif ext == '.txt':
            return file_path.read_text(encoding="utf-8")
    except Exception as e:
        logging.error(f"Error extracting text: {str(e)}")
        return None
    return ""

def generate_mcqs_from_text(text, num_questions):
    prompt = f"""
    Generate {num_questions} multiple-choice questions (MCQs) from the following text:
    
    - Ensure questions are **clear and relevant**.
    - Provide **4 distinct answer choices**.
    - Indicate the **correct answer** at the end.
    
    TEXT:
    '{text}'
    
    **Output Format:**
    ## MCQ
    Question: [question]
    A) [option A]
    B) [option B]
    C) [option C]
    D) [option D]
    Correct Answer: [correct option]
    """
    try:
        response = model.generate_content(prompt)
        return response.text.strip() if response else "No MCQs generated."
    except Exception as e:
        logging.error(f"Error generating MCQs: {str(e)}")
        return None

def save_text_file(content, filename):
    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    file_path = RESULTS_FOLDER / filename
    content_with_timestamp = f"Generated on: {timestamp}\n\n{content}"
    try:
        file_path.write_text(content_with_timestamp, encoding="utf-8")
        return file_path
    except Exception as e:
        logging.error(f"Error saving text file: {str(e)}")
        return None

def generate_pdf(mcqs, filename):
    try:
        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=10)
        pdf.add_page()

        pdf.set_font("Arial", "B", 16)
        pdf.cell(200, 10, "Generated MCQs".encode('latin-1', 'replace').decode('latin-1'), ln=True, align="C")
        pdf.ln(5)

        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        pdf.set_font("Arial", "I", 10)
        pdf.cell(200, 10, f"Generated on: {timestamp}".encode('latin-1', 'replace').decode('latin-1'), ln=True, align="C")
        pdf.ln(10)

        pdf.set_font("Arial", size=12)

        for i, mcq in enumerate(mcqs.split("## MCQ"), start=0):
            if mcq.strip():
                lines = mcq.strip().split("\n")
                correct_answer = next((line.split(":")[-1].strip() for line in lines if "Correct Answer:" in line), None)

                pdf.set_font("Arial", "B", 12)
                pdf.cell(0, 10, f"Q{i}: {lines[0].split(':')[-1].strip()}".encode('latin-1', 'replace').decode('latin-1'), ln=True)

                pdf.set_font("Arial", size=12)
                for option in lines[1:]:
                    if option.startswith("Correct Answer:"):
                        continue
                    
                    if correct_answer and option.strip().startswith(correct_answer + ")"):
                        pdf.set_text_color(34, 139, 34)  # Green for correct answer
                        pdf.set_font("Arial", "B", 12)  # Bold for correct answer
                    else:
                        pdf.set_text_color(0, 0, 0)  # Black for incorrect options
                        pdf.set_font("Arial", size=12)

                    pdf.cell(0, 10, option.strip().encode('latin-1', 'replace').decode('latin-1'), ln=True)

                pdf.set_text_color(0, 128, 0)  # Green for correct answer
                pdf.cell(0, 10, f"Correct Answer: {correct_answer}".encode('latin-1', 'replace').decode('latin-1'), ln=True)
                pdf.set_text_color(0, 0, 0)
                pdf.ln(5)

        pdf_path = RESULTS_FOLDER / filename
        pdf.output(pdf_path)
        return pdf_path
    except Exception as e:
        logging.error(f"Error generating PDF: {str(e)}")
        return None

def generate_docx(mcqs, filename):
    try:
        doc = Document()
        
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        doc.add_paragraph(f"Generated on: {timestamp}\n").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        for i, mcq in enumerate(mcqs.split("## MCQ"), start=0):
            if mcq.strip():
                lines = mcq.strip().split("\n")
                correct_answer = next((line.split(":")[-1].strip() for line in lines if "Correct Answer:" in line), None)

                para = doc.add_paragraph(f"Q{i}: {lines[0].split(':')[-1].strip()}")
                para.runs[0].bold = True

                for option in lines[1:]:
                    if option.startswith("Correct Answer:"):
                        continue

                    para = doc.add_paragraph(option.strip())
                    if correct_answer and option.strip().startswith(correct_answer + ")"):
                        for run in para.runs:
                            run.font.color.rgb = RGBColor(34, 139, 34)  # Green for correct answer

                para = doc.add_paragraph(f"Correct Answer: {correct_answer}")
                for run in para.runs:
                    run.font.color.rgb = RGBColor(34, 139, 34)  # Green for correct answer

                doc.add_paragraph("\n")

        docx_path = RESULTS_FOLDER / filename
        doc.save(docx_path)
        return docx_path
    except Exception as e:
        logging.error(f"Error generating DOCX: {str(e)}")
        return None

# ============================
# 📌 Flask Routes
# ============================
@app.route('/')
def index():
    return render_template('index.html')

@app.route('/generate', methods=['POST'])
def generate_mcqs():
    if 'file' not in request.files:
        flash("No file uploaded", "error")
        return redirect(request.url)

    file = request.files['file']
    if file.filename == '':
        flash("No file selected", "error")
        return redirect(request.url)

    if not allowed_file(file.filename):
        flash("Invalid file format. Allowed formats: PDF, TXT, DOCX", "error")
        return redirect(request.url)

    filename = secure_filename(file.filename)
    file_path = UPLOAD_FOLDER / filename
    file.save(file_path)

    text = extract_text_from_file(file_path)
    if not text:
        flash("Could not extract text from file", "error")
        return redirect(request.url)

    try:
        num_questions = int(request.form['num_questions'])
        if num_questions <= 0:
            flash("Number of questions must be greater than zero", "error")
            return redirect(request.url)
    except ValueError:
        flash("Invalid number of questions", "error")
        return redirect(request.url)

    mcqs = generate_mcqs_from_text(text, num_questions)
    if not mcqs:
        flash("Failed to generate MCQs", "error")
        return redirect(request.url)

    base_filename = filename.rsplit('.', 1)[0]
    txt_filename = f"generated_mcqs_{base_filename}.txt"
    pdf_filename = f"generated_mcqs_{base_filename}.pdf"
    docx_filename = f"generated_mcqs_{base_filename}.docx"

    if not save_text_file(mcqs, txt_filename):
        flash("Failed to save text file", "error")
        return redirect(request.url)

    if not generate_pdf(mcqs, pdf_filename):
        flash("Failed to generate PDF", "error")
        return redirect(request.url)

    if not generate_docx(mcqs, docx_filename):
        flash("Failed to generate DOCX", "error")
        return redirect(request.url)

    return render_template('results.html', mcqs=mcqs, txt_filename=txt_filename, pdf_filename=pdf_filename, docx_filename=docx_filename)

@app.route('/download/<filename>')
def download_file(filename):
    file_path = RESULTS_FOLDER / filename
    if not file_path.exists():
        flash("File not found", "error")
        return redirect('/')
    return send_file(file_path, as_attachment=True)

# ============================
# 🚀 Run the Flask App
# ============================
if __name__ == "__main__":
    logging.basicConfig(level=logging.INFO)
    app.run(host="0.0.0.0", port=5000, debug=True, threaded=True)  